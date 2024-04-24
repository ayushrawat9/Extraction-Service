from celery import shared_task
from text_extraction.celery import app
from django.core.mail import send_mail
import PyPDF2
import docx
from urllib.parse import urlsplit
from io import BytesIO
import base64
from .models import ExtractedText
from django.contrib.auth.models import User

@shared_task(bind=True)
def send_email_task(filename, text, subject='', message='', recipient_list=['ayush.rawat.1357@gmail.com']):
    subject = '''Text Extraction Completed'''
    message = f'''
        File Name : {filename}
        Extracted Text : {text}
    '''
    try:
        send_mail(subject, message, 'mailtrap@demomailtrap.com', recipient_list)
        print('Mail Sent Successfully!')
    except Exception as e:
        print('Mail Error occurred')
        print(e)
    
def save_res(text, filename, user_id):
    if text:
        obj = ExtractedText()
        obj.extracted_text = text
        obj.file_name = filename
        obj.created_by = User.objects.get(pk=user_id)
        obj.save()
        send_email_task.delay(filename, text)

@shared_task(bind=True)
def extract_text_from_pdf(self, filename, uploaded_file, user_id):
    try:
        uploaded_file = base64.b64decode(uploaded_file)
        if isinstance(uploaded_file, bytes):
            uploaded_file = BytesIO(uploaded_file)
        reader = PyPDF2.PdfReader(uploaded_file)
        pdf_text = ''
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            pdf_text += page.extract_text()
        # return pdf_text
        save_res(pdf_text, filename, user_id)
        return pdf_text
    except Exception as e:
        return f"Error: {e}"

@shared_task(bind=True)
def extract_text_from_docx_or_doc(filename, uploaded_file, file_extension, user_id):
    try:
        uploaded_file = base64.b64decode(uploaded_file)
        if isinstance(uploaded_file, bytes):
            uploaded_file = BytesIO(uploaded_file)
        if file_extension == 'docx':
            docx_text = docx.Document(uploaded_file)
        else:
            docx_text = docx.Document(uploaded_file)
        full_text = []
        for paragraph in docx_text.paragraphs:
            full_text.append(paragraph.text)
        # return '\n'.join(full_text)
        save_res('\n'.join(full_text), filename, user_id)
    except Exception as e:
        return f"Error: {e}"